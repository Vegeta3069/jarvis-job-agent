#!/usr/bin/env python3
"""
Jarvis MCP Server v2 — H1B Job Hunter
Trigger: "wake up Jarvis"

v2 changes:
  - Sourcing: ATS feeds only (jarvis_sourcing.py + sponsors.yaml).
    DDGS / web search REMOVED. No padding: short day = short list.
  - Freshness: <= 14 days, from real ATS timestamps. Unknown date = rejected.
  - Every link liveness-checked before delivery; carryover re-checked daily.
  - Carryover: unapplied jobs roll into today's list until applied or closed.
  - Tracker migration: adds posted_at/ats/employer_type/apply_mode/resume_file,
    archives legacy junk-source rows (linkedin/dice/bing/etc) once.
  - Numbering: ONE absolute tracker row number everywhere (digest, list_jobs,
    mark_applied, tailor_resume). v1 mismatch bug fixed.
  - tailor_resume(n): tailored resume DOCX per job from profile/resume_base.docx
    via Anthropic API. Re-emphasis of REAL experience only; fabrication banned.
"""

import csv
import json
import os
import shutil
from datetime import date, datetime, timezone
from pathlib import Path

import yaml
from mcp.server.fastmcp import FastMCP

import jarvis_sourcing as src
import web_sourcing as web

BASE_DIR = Path(__file__).parent
JOBS_DIR = BASE_DIR / "jobs"
RESUME_DIR = BASE_DIR / "resumes"
PROFILE_DIR = BASE_DIR / "profile"
TRACKER_FILE = BASE_DIR / "tracker.csv"
SPONSORS_FILE = BASE_DIR / "sponsors.yaml"
BASE_RESUME = PROFILE_DIR / "resume_base.docx"
PROFILE_FILE = BASE_DIR / "profile.yaml"

DAILY_CAP = 60          # standalone find_jobs / search_web_jobs ceiling
DAILY_TARGET = 30       # combined daily_jobs list size (sponsors first, then web)

FIELDS = ["date_found", "title", "company", "location", "link", "status",
          "date_applied", "notes",
          "posted_at", "ats", "employer_type", "apply_mode", "resume_file"]

mcp = FastMCP("Jarvis")

TODAY = lambda: date.today().strftime("%Y-%m-%d")


# ── tracker I/O + one-time migration ─────────────────────────────────────────
def _ensure_setup():
    JOBS_DIR.mkdir(exist_ok=True)
    RESUME_DIR.mkdir(exist_ok=True)
    PROFILE_DIR.mkdir(exist_ok=True)
    if not TRACKER_FILE.exists():
        with open(TRACKER_FILE, "w", newline="") as f:
            csv.writer(f).writerow(FIELDS)


def _load_tracker() -> list[dict]:
    """Load rows; migrate schema and archive legacy junk sources exactly once."""
    _ensure_setup()
    with open(TRACKER_FILE, newline="") as f:
        rows = list(csv.DictReader(f))

    migrated = False
    for r in rows:
        for fld in FIELDS:
            if r.get(fld) is None:
                r[fld] = ""
                migrated = True
        # one-time archive of UNTAGGED legacy rows from banned sources
        # (linkedin/dice/bing/...). Rows with an `ats` tag — including
        # web-search finds (ats="web") — are legitimate and never touched.
        if (r["status"] == "not_applied" and not r.get("ats")
                and not src.allowed_domain(r["link"])):
            r["status"] = "archived"
            r["notes"] = (r["notes"] + " | " if r["notes"] else "") + "legacy junk source"
            migrated = True

    if migrated:
        backup = TRACKER_FILE.with_suffix(".csv.bak")
        if not backup.exists():
            shutil.copy(TRACKER_FILE, backup)
        _save_tracker(rows)
    return rows


def _save_tracker(rows: list[dict]):
    with open(TRACKER_FILE, "w", newline="") as f:
        w = csv.DictWriter(f, fieldnames=FIELDS, extrasaction="ignore")
        w.writeheader()
        w.writerows(rows)


def _fmt_row(n: int, r: dict) -> list[str]:
    badge = "🤖 auto" if r.get("apply_mode") == "auto" else "🖐 manual"
    tier = "🏢" if r.get("employer_type", "direct") == "direct" else "🏗 consultancy"
    out = [f"**#{n}. {r['title']}** — {r['company']}",
           f"   📍 {r['location']} | 📅 posted {r.get('posted_at') or '?'} | {badge} | {tier}",
           f"   🔗 {r['link']}"]
    if r.get("resume_file"):
        out.append(f"   📄 tailored resume: `{r['resume_file']}`")
    return out + [""]


# ── MCP tools ────────────────────────────────────────────────────────────────
@mcp.tool()
def daily_jobs() -> str:
    """
    THE daily run — call this when the user says 'wake up Jarvis'. Delivers up to
    30 BRAND-NEW jobs each day, sponsor list first then web:
      1) sponsor list (sponsors.yaml ATS feeds) fills the 30 first
      2) the web (Adzuna) fills whatever slots remain up to 30
    Carryover (still-open unapplied jobs from earlier days) is re-verified and
    shown SEPARATELY below the new list — it does NOT consume the 30 slots.
    Every job matches your résumé profile, was posted within max_age_days, is
    US-based, de-duplicated, and has a verified-live apply link. New finds are
    added to the tracker; a digest is written to jobs/<date>_daily.md.
    """
    rows = _load_tracker()
    profile = src.load_profile()
    sess = src.new_session()
    today = TODAY()

    # 1) carryover: re-gate against the profile, then liveness-check
    carry_idx = [i for i, r in enumerate(rows) if r["status"] == "not_applied"]
    regated = 0
    survivors = []
    for i in carry_idx:
        r = rows[i]
        if src.passes_content_gates(r["title"], r["location"], profile):
            survivors.append(i)
        else:
            r["status"] = "archived"
            r["notes"] = (r["notes"] + " | " if r["notes"] else "") + \
                         f"fails gates, archived {today}"
            regated += 1
    carry_objs = {i: src.Job(rows[i]["company"], rows[i]["title"], rows[i]["location"],
                             rows[i]["link"], rows[i].get("ats", ""),
                             rows[i].get("employer_type", "direct"), None,
                             rows[i].get("apply_mode", "manual"))
                  for i in survivors}
    alive = {j.url for j in src.verify_live(sess, list(carry_objs.values()))}
    closed = 0
    for i in survivors:
        if rows[i]["link"] not in alive:
            rows[i]["status"] = "closed"
            rows[i]["notes"] = (rows[i]["notes"] + " | " if rows[i]["notes"] else "") + \
                               f"req closed, detected {today}"
            closed += 1
    carry_live = [i for i in survivors if rows[i]["link"] in alive]

    known = {src.canonical(r["link"]) for r in rows if r.get("link")}
    room = DAILY_TARGET          # 30 brand-new; carryover is shown separately, not counted

    # 2) sponsor list first
    try:
        cfg = src.load_config(str(SPONSORS_FILE))
    except FileNotFoundError:
        cfg = {"companies": []}
    raw, sstats = src.fetch_all(cfg, sess) if cfg.get("companies") else \
        ([], {"sources_ok": 0, "sources_failed": 0, "failed_names": []})
    gated, gstats = src.apply_gates(raw, known, profile)
    ats_live = src.verify_live(sess, gated)
    ats_dead = len(gated) - len(ats_live)
    ats_fresh = src.rank(ats_live)
    ats_take = ats_fresh[:room]
    for j in ats_take:
        known.add(src.canonical(j.url))

    # 3) web fills any remaining slots
    room2 = room - len(ats_take)
    web_take, wstats, web_status = [], None, "not needed"
    if room2 > 0:
        try:
            cands = web.search_adzuna(profile, max_days=profile["max_age_days"])
            web_deliv, wstats = web.verify_and_gate(cands, known, profile, sess)
            web_take = web_deliv[:room2]
            web_status = "ok"
        except RuntimeError:
            web_status = "skipped (no ADZUNA key)"

    # append new finds (sponsors then web)
    new_start = len(rows)
    for j in ats_take:
        rows.append({
            "date_found": today, "title": j.title, "company": j.company,
            "location": j.location, "link": j.url, "status": "not_applied",
            "date_applied": "", "notes": "",
            "posted_at": j.posted_at.strftime("%Y-%m-%d") if j.posted_at else "",
            "ats": j.ats, "employer_type": j.employer_type,
            "apply_mode": j.apply_mode, "resume_file": "",
        })
    for c in web_take:
        rows.append({
            "date_found": today, "title": c["title"], "company": c["company"],
            "location": c["location"], "link": c["url"], "status": "not_applied",
            "date_applied": "", "notes": "web search",
            "posted_at": c["posted"].strftime("%Y-%m-%d") if c["posted"] else "",
            "ats": "web", "employer_type": "direct",
            "apply_mode": "manual", "resume_file": "",
        })
    _save_tracker(rows)

    new_count = len(ats_take) + len(web_take)
    role_note = "" if PROFILE_FILE.exists() else \
        "  _(default filter — run setup_profile with your resume to personalize)_"
    stat_lines = [
        f"| Sponsors: sources ok/fail | {sstats['sources_ok']} / {sstats['sources_failed']} |",
        f"| Sponsors: pulled / matched / dead-link | {len(raw)} / {len(ats_fresh)} / {ats_dead} |",
        f"| Web (Adzuna): {web_status} | "
        f"{(str(wstats['fetched'])+' fetched, '+str(len(web_take))+' kept') if wstats else '—'} |",
        f"| **New today (target {DAILY_TARGET})** | "
        f"**{new_count}**  ({len(ats_take)} sponsors + {len(web_take)} web) |",
        f"| Carryover still open / closed / archived | {len(carry_live)} / {closed} / {regated} |",
    ]
    lines = [f"# Jarvis Daily — {today}", "",
             f"**Target role:** {profile.get('target_role', '—')}{role_note}", "",
             "| Pipeline | Count |", "|---|---|", *stat_lines, "", "---", ""]
    if ats_take:
        lines.append(f"## 🏢 New from your sponsor list ({len(ats_take)})\n")
        for k in range(len(ats_take)):
            lines += _fmt_row(new_start + k + 1, rows[new_start + k])
    if web_take:
        lines.append(f"## 🌐 New from the web — every link verified ({len(web_take)})\n")
        base = new_start + len(ats_take)
        for k in range(len(web_take)):
            lines += _fmt_row(base + k + 1, rows[base + k])
    if new_count == 0:
        lines.append("No new matches survived verification today. No padding.")
    if carry_live:
        lines.append(f"\n---\n\n## 🔁 Carryover — still open from earlier days "
                     f"({len(carry_live)}, not counted in the {DAILY_TARGET})\n")
        for i in carry_live:
            lines += _fmt_row(i + 1, rows[i])

    lines += ["---",
              f"💬 'mark N applied' · 'tailor resume N' · "
              f"{new_count}/{DAILY_TARGET} new today · {len(carry_live)} carried over"]
    digest = "\n".join(lines)
    with open(JOBS_DIR / f"{today}_daily.md", "w") as f:
        f.write(digest + "\n")
    return digest


@mcp.tool()
def find_jobs() -> str:
    """
    Sponsor-list-only run (for the full daily run use daily_jobs, which also adds
    web results). Searches just your sponsors.yaml ATS feeds, capped at DAILY_CAP.
    1) Re-verifies every carried-over unapplied job (closed reqs marked closed)
    2) Pulls fresh postings from sponsor companies' own ATS feeds (sponsors.yaml)
    3) Gates: domain, dedup, freshness, title, US location, eligibility. The
       TITLE filter is the one derived from your resume (profile.yaml); run
       setup_profile first to personalize, otherwise a default filter is used.
    4) Liveness-checks every new link
    5) Delivers carryover + new, capped at DAILY_CAP, never padded, with honest
       rejection stats. Writes digest to jobs/<date>_jobs.md and appends
       new rows to the tracker.
    """
    rows = _load_tracker()
    profile = src.load_profile()
    sess = src.new_session()
    today = TODAY()

    # 1) carryover: every unapplied row, re-screened then re-verified
    carry_idx = [i for i, r in enumerate(rows) if r["status"] == "not_applied"]

    # 1a) re-gate content (title/location/eligibility) so legacy rows that no
    #     longer match — intern/garbage titles, non-US — drop out of carryover.
    regated = 0
    survivors = []
    for i in carry_idx:
        r = rows[i]
        if src.passes_content_gates(r["title"], r["location"], profile):
            survivors.append(i)
        else:
            r["status"] = "archived"
            r["notes"] = (r["notes"] + " | " if r["notes"] else "") + \
                         f"fails v2 gates, archived {today}"
            regated += 1

    # 1b) liveness: closed reqs marked closed
    carry_jobs = {i: src.Job(rows[i]["company"], rows[i]["title"], rows[i]["location"],
                             rows[i]["link"], rows[i].get("ats", ""),
                             rows[i].get("employer_type", "direct"), None,
                             rows[i].get("apply_mode", "manual"))
                  for i in survivors}
    alive_urls = {j.url for j in src.verify_live(sess, list(carry_jobs.values()))}
    closed = 0
    for i in survivors:
        if rows[i]["link"] not in alive_urls:
            rows[i]["status"] = "closed"
            rows[i]["notes"] = (rows[i]["notes"] + " | " if rows[i]["notes"] else "") + \
                               f"req closed, detected {today}"
            closed += 1
    carry_live = [i for i in survivors if rows[i]["link"] in alive_urls]

    # 2) fresh pulls
    try:
        cfg = src.load_config(str(SPONSORS_FILE))
    except FileNotFoundError:
        return ("⛔ sponsors.yaml not found next to jarvis_mcp.py. "
                "That file is the company registry; sourcing cannot run without it.")
    raw, sstats = src.fetch_all(cfg, sess)

    # 3) gates + 4) liveness
    known = {src.canonical(r["link"]) for r in rows}
    gated, gstats = src.apply_gates(raw, known, profile)
    fresh = src.verify_live(sess, gated)
    dead = len(gated) - len(fresh)
    fresh = src.rank(fresh)

    # 5) cap, never pad
    room = max(0, DAILY_CAP - len(carry_live))
    fresh = fresh[:room]

    new_start = len(rows)
    for j in fresh:
        rows.append({
            "date_found": today, "title": j.title, "company": j.company,
            "location": j.location, "link": j.url, "status": "not_applied",
            "date_applied": "", "notes": "",
            "posted_at": j.posted_at.strftime("%Y-%m-%d") if j.posted_at else "",
            "ats": j.ats, "employer_type": j.employer_type,
            "apply_mode": j.apply_mode, "resume_file": "",
        })
    _save_tracker(rows)

    # digest
    stat_lines = [
        f"| Sources OK / failed | {sstats['sources_ok']} / {sstats['sources_failed']} |",
        f"| Raw postings pulled | {len(raw)} |",
        f"| Rejected: off-allowlist domain | {gstats['rej_domain']} |",
        f"| Rejected: duplicate | {gstats['rej_duplicate']} |",
        f"| Rejected: older than {profile['max_age_days']}d / no date | {gstats['rej_stale']} |",
        f"| Rejected: title mismatch | {gstats['rej_title']} |",
        f"| Rejected: non-US / no location | {gstats['rej_location']} |",
        f"| Rejected: clearance/eligibility | {gstats['rej_eligibility']} |",
        f"| Rejected: dead link | {dead} |",
        f"| Carryover alive / closed overnight | {len(carry_live)} / {closed} |",
        f"| Carryover archived (failed v2 gates) | {regated} |",
        f"| **Delivered today** | **{len(carry_live) + len(fresh)}** |",
    ]
    if sstats["failed_names"]:
        stat_lines.append(f"| Failed sources | {', '.join(sstats['failed_names'])} |")

    role_note = "" if PROFILE_FILE.exists() else \
        "  _(default filter — run `setup_profile` with your resume to personalize)_"
    lines = [f"# Jarvis Digest — {today}", "",
             f"**Target role:** {profile.get('target_role', '—')}{role_note}", "",
             "| Pipeline | Count |", "|---|---|", *stat_lines, "", "---", ""]

    if carry_live:
        lines.append(f"## 🔁 Carryover — still open, still unapplied ({len(carry_live)})\n")
        for i in carry_live:
            lines += _fmt_row(i + 1, rows[i])
    if fresh:
        lines.append(f"## 🆕 New today ({len(fresh)})\n")
        for k in range(len(fresh)):
            lines += _fmt_row(new_start + k + 1, rows[new_start + k])
    if not carry_live and not fresh:
        lines.append("Nothing credible survived the gates today. "
                     "That is the honest answer; no padding.")

    lines += ["---",
              "💬 'mark N applied' to log an application | "
              "'tailor resume N' for a job-specific resume | "
              "numbers are permanent tracker row IDs"]

    digest = "\n".join(lines)
    with open(JOBS_DIR / f"{today}_jobs.md", "w") as f:
        f.write(digest + "\n")
    return digest


@mcp.tool()
def search_web_jobs() -> str:
    """
    Internet-wide job search — SEPARATE from the ATS find_jobs. Scans jobs across
    the web via the Adzuna API, keeps only those that match YOUR RESUME PROFILE,
    were posted within max_age_days, are US-based, and — critically —
    liveness-checks EVERY apply link before delivering (dead/false links are
    dropped, not shown). New finds are added to the tracker tagged source 'web'.
    Requires ADZUNA_APP_ID + ADZUNA_APP_KEY in the environment (free key at
    https://developer.adzuna.com). Trigger: 'search the web' / 'wide search'.
    """
    rows = _load_tracker()
    profile = src.load_profile()
    today = TODAY()

    try:
        cands = web.search_adzuna(profile, max_days=profile["max_age_days"])
    except RuntimeError as e:
        return (f"⛔ {e}.\nGet a free key at https://developer.adzuna.com and add it to "
                "claude_desktop_config.json:\n"
                '    "env": { "ADZUNA_APP_ID": "...", "ADZUNA_APP_KEY": "..." }')

    prior = {src.canonical(r["link"]) for r in rows if r.get("link")}
    sess = src.new_session()
    delivered, st = web.verify_and_gate(cands, prior, profile, sess)
    delivered = delivered[:DAILY_CAP]

    new_start = len(rows)
    for c in delivered:
        rows.append({
            "date_found": today, "title": c["title"], "company": c["company"],
            "location": c["location"], "link": c["url"], "status": "not_applied",
            "date_applied": "", "notes": "web search",
            "posted_at": c["posted"].strftime("%Y-%m-%d") if c["posted"] else "",
            "ats": "web", "employer_type": "direct",
            "apply_mode": "manual", "resume_file": "",
        })
    _save_tracker(rows)

    role_note = "" if PROFILE_FILE.exists() else \
        "  _(default filter — run setup_profile with your resume to personalize)_"
    stat_lines = [
        f"| Fetched from web (Adzuna) | {st['fetched']} |",
        f"| Rejected: older than {profile['max_age_days']}d / no date | {st['rej_stale']} |",
        f"| Rejected: title mismatch | {st['rej_title']} |",
        f"| Rejected: non-US | {st['rej_location']} |",
        f"| Rejected: clearance/eligibility | {st['rej_eligibility']} |",
        f"| Rejected: duplicate (vs tracker/batch) | {st['rej_duplicate']} |",
        f"| Rejected: DEAD LINK (failed verification) | {st['rej_dead']} |",
        f"| **Delivered (every link verified live)** | **{len(delivered)}** |",
    ]
    lines = [f"# Jarvis Web Search — {today}", "",
             f"**Target role:** {profile.get('target_role', '—')}{role_note}", "",
             "| Pipeline | Count |", "|---|---|", *stat_lines, "", "---", ""]
    if delivered:
        lines.append(f"## 🌐 New from the web ({len(delivered)}) — every link checked live\n")
        for k in range(len(delivered)):
            lines += _fmt_row(new_start + k + 1, rows[new_start + k])
    else:
        lines.append("Nothing survived verification today. No dead links, no padding.")

    digest = "\n".join(lines)
    with open(JOBS_DIR / f"{today}_web_jobs.md", "w") as f:
        f.write(digest + "\n")
    return digest


@mcp.tool()
def list_jobs(date_filter: str = "today") -> str:
    """
    List tracker jobs. date_filter: 'today', 'all', 'pending' (all unapplied),
    or a date like '2026-05-31'. Numbers are permanent tracker row IDs and
    match mark_applied / tailor_resume.
    """
    rows = _load_tracker()
    if not rows:
        return "Tracker empty. Say 'wake up Jarvis' to run the daily search."

    if date_filter == "pending":
        sel = [(i, r) for i, r in enumerate(rows) if r["status"] == "not_applied"]
    elif date_filter == "all":
        sel = list(enumerate(rows))
    else:
        target = TODAY() if date_filter == "today" else date_filter
        sel = [(i, r) for i, r in enumerate(rows) if r["date_found"] == target]

    if not sel:
        return f"No rows for '{date_filter}'."

    pend = [(i, r) for i, r in sel if r["status"] == "not_applied"]
    appl = [(i, r) for i, r in sel if r["status"] == "applied"]
    other = len(sel) - len(pend) - len(appl)

    lines = [f"📋 **{date_filter}** — {len(sel)} rows | {len(appl)} applied | "
             f"{len(pend)} pending | {other} closed/archived\n"]
    if pend:
        lines.append("### 🔵 Pending\n")
        for i, r in pend:
            lines += _fmt_row(i + 1, r)
    if appl:
        lines.append("### ✅ Applied\n")
        for i, r in appl:
            lines.append(f"#{i + 1} ~~{r['title']}~~ — {r['company']} "
                         f"(applied {r['date_applied']})")
    return "\n".join(lines)


@mcp.tool()
def mark_applied(job_number: int, notes: str = "") -> str:
    """
    Mark a job applied by its permanent tracker row number (the #N shown in
    digests and list_jobs). Example: mark_applied(112, 'via Greenhouse')
    """
    rows = _load_tracker()
    idx = job_number - 1
    if idx < 0 or idx >= len(rows):
        return f"Invalid row #{job_number}. Valid: 1–{len(rows)}."
    r = rows[idx]
    if r["status"] == "applied":
        return f"#{job_number} already applied on {r['date_applied']}: {r['title']} @ {r['company']}."
    if r["status"] in ("closed", "archived"):
        return (f"#{job_number} is {r['status']} ({r['title']} @ {r['company']}). "
                f"If you genuinely applied, say so and I will override.")
    r["status"] = "applied"
    r["date_applied"] = TODAY()
    if notes:
        r["notes"] = (r["notes"] + " | " if r["notes"] else "") + notes
    _save_tracker(rows)
    return f"✅ #{job_number} applied: **{r['title']}** @ {r['company']} ({TODAY()})"


@mcp.tool()
def get_stats() -> str:
    """Application stats: totals, per-day found/applied, pipeline health."""
    rows = _load_tracker()
    if not rows:
        return "No jobs tracked yet."
    n = lambda s: sum(1 for r in rows if r["status"] == s)
    total, applied = len(rows), n("applied")
    lines = ["## 📊 Jarvis Stats", "", "| Metric | Count |", "|---|---|",
             f"| Total tracked | {total} |",
             f"| Applied | {applied} |",
             f"| Pending | {n('not_applied')} |",
             f"| Closed (req gone) | {n('closed')} |",
             f"| Archived (legacy junk) | {n('archived')} |",
             f"| Application rate (excl. junk) | "
             f"{int(applied / max(1, total - n('archived')) * 100)}% |", ""]
    per_day: dict[str, list[int]] = {}
    for r in rows:
        d = per_day.setdefault(r["date_found"], [0, 0])
        d[0] += 1
        if r["status"] == "applied":
            d[1] += 1
    lines.append("### Per day (found | applied)")
    for d in sorted(per_day):
        lines.append(f"- **{d}**: {per_day[d][0]} | {per_day[d][1]}")
    return "\n".join(lines)


@mcp.tool()
def open_digest(target_date: str = "") -> str:
    """Path to a day's digest markdown. Optional date '2026-06-10'."""
    d = target_date or TODAY()
    p = JOBS_DIR / f"{d}_jobs.md"
    if not p.exists():
        return f"No digest for {d}. Say 'wake up Jarvis' to generate today's."
    return f"📄 `{p}`\n\nTerminal: `open '{p}'`"


# ── resume → role profile ────────────────────────────────────────────────────
PROFILE_SYSTEM = (
    "You convert a resume into a job-search title filter. Output ONLY minified "
    "JSON (no prose, no code fences) with exactly these keys: "
    '{"target_role": "<short label, e.g. Senior Data Engineer>", '
    '"title_include": ["lowercase substrings that SHOULD appear in a matching '
    'job title"], "title_exclude": ["lowercase substrings that DISQUALIFY a '
    'title"]}. '
    "Base it on the candidate's real domain and seniority. title_include = 4-8 "
    "short lowercase keywords/phrases for the roles they should target. "
    "title_exclude = wrong levels/domains to avoid (e.g. for a senior IC: "
    "'intern','new grad','manager','director'). Keywords are matched by substring "
    "against the lowercased job title, so keep them short and generic."
)


@mcp.tool()
def setup_profile() -> str:
    """
    Personalize Jarvis to a resume — THIS is the main input. Reads
    profile/resume_base.docx, derives the target-role title filter via the
    Anthropic API, and writes profile.yaml. Run once after dropping in a resume
    (or after updating it); find_jobs then matches roles to that resume against
    the same sponsor company list. Requires ANTHROPIC_API_KEY in the environment.
    """
    import requests as rq

    _ensure_setup()
    if not BASE_RESUME.exists():
        return ("⛔ No resume found. Drop your resume at "
                f"`{BASE_RESUME}` (exact filename resume_base.docx), then run setup_profile.")
    key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not key:
        return "⛔ ANTHROPIC_API_KEY not set in Jarvis's environment (claude_desktop_config.json)."

    text = _docx_text(BASE_RESUME)
    resp = rq.post("https://api.anthropic.com/v1/messages",
                   headers={"x-api-key": key, "anthropic-version": "2023-06-01",
                            "content-type": "application/json"},
                   json={"model": "claude-sonnet-4-5", "max_tokens": 600,
                         "system": PROFILE_SYSTEM,
                         "messages": [{"role": "user", "content": f"RESUME:\n{text[:8000]}"}]},
                   timeout=60)
    if resp.status_code != 200:
        return f"⛔ Anthropic API error {resp.status_code}: {resp.text[:200]}"

    raw = "".join(b.get("text", "") for b in resp.json().get("content", [])).strip()
    raw = raw.strip("`")
    if raw.lower().startswith("json"):
        raw = raw[4:]
    try:
        derived = json.loads(raw)
        inc = [s.lower() for s in derived["title_include"] if s and s.strip()]
        exc = [s.lower() for s in derived.get("title_exclude", []) if s and s.strip()]
        role = (derived.get("target_role") or "").strip() or "—"
        assert inc
    except Exception as e:
        return f"⛔ Could not parse role from resume ({e}). Model returned:\n{raw[:300]}"

    # Resume drives ROLE only; eligibility + freshness stay at safe defaults.
    profile = {
        "target_role": role,
        "title_include": inc,
        "title_exclude": exc or src.DEFAULT_PROFILE["title_exclude"],
        "eligibility_exclude": src.DEFAULT_PROFILE["eligibility_exclude"],
        "max_age_days": src.DEFAULT_PROFILE["max_age_days"],
    }
    with open(PROFILE_FILE, "w") as f:
        yaml.safe_dump(profile, f, sort_keys=False, default_flow_style=False)
    return ("✅ Profile set from your resume → `profile.yaml`\n"
            f"🎯 Target role: **{role}**\n"
            f"🔎 Matching titles containing: {', '.join(inc)}\n"
            f"🚫 Excluding: {', '.join(profile['title_exclude'])}\n\n"
            "Now say **'wake up Jarvis'** to search the sponsor list for these roles.")


# ── resume tailoring ─────────────────────────────────────────────────────────
TAILOR_SYSTEM = (
    "You tailor a resume to a job posting. HARD RULES: use ONLY experience, "
    "skills, employers, titles, and dates present in the base resume. You may "
    "reorder bullets, reweight emphasis, mirror the posting's terminology for "
    "skills the candidate genuinely has, and rewrite the summary. You may NOT "
    "invent, exaggerate, add tools the base resume lacks, or alter dates/titles. "
    "The candidate is an H-1B holder; fabrication carries immigration risk. "
    "Return ONLY the tailored resume as plain structured text with section "
    "headers in ALL CAPS, no commentary, no markdown fences."
)


def _docx_text(path: Path) -> str:
    from docx import Document
    return "\n".join(p.text for p in Document(str(path)).paragraphs if p.text.strip())


def _text_to_docx(text: str, out: Path):
    from docx import Document
    doc = Document()
    for line in text.splitlines():
        s = line.strip()
        if not s:
            continue
        if s.isupper() and len(s) < 60:
            doc.add_heading(s.title(), level=2)
        elif s.startswith(("-", "•", "*")):
            doc.add_paragraph(s.lstrip("-•* ").strip(), style="List Bullet")
        else:
            doc.add_paragraph(s)
    doc.save(str(out))


@mcp.tool()
def tailor_resume(job_number: int) -> str:
    """
    Generate a job-tailored resume DOCX for tracker row #N and link it in the
    tracker's resume_file column. Requires profile/resume_base.docx and
    ANTHROPIC_API_KEY in the environment. Re-emphasizes real experience only;
    never fabricates. Export to PDF from Word/Pages before applying.
    """
    import requests as rq

    rows = _load_tracker()
    idx = job_number - 1
    if idx < 0 or idx >= len(rows):
        return f"Invalid row #{job_number}."
    r = rows[idx]
    if not BASE_RESUME.exists():
        return ("⛔ profile/resume_base.docx missing. Drop your current resume "
                "there (exact filename) and retry.")
    key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not key:
        return "⛔ ANTHROPIC_API_KEY not set in Jarvis's environment."

    # posting text: greenhouse content API when possible, else page fetch
    posting = f"{r['title']} at {r['company']} ({r['location']})"
    try:
        sess = src.new_session()
        if r.get("ats") == "greenhouse" and "/jobs/" in r["link"]:
            import re as _re
            m = _re.search(r"boards.greenhouse.io/([^/]+)/jobs/(\d+)", r["link"]) or \
                _re.search(r"greenhouse.io/([^/]+)/jobs/(\d+)", r["link"])
            if m:
                resp = sess.get(f"https://boards-api.greenhouse.io/v1/boards/"
                                f"{m.group(1)}/jobs/{m.group(2)}", timeout=15).json()
                import html, re as _re2
                posting = html.unescape(_re2.sub(r"<[^>]+>", " ", resp.get("content", "")))[:6000]
        else:
            posting = sess.get(r["link"], timeout=15).text[:6000]
    except Exception:
        pass  # fall back to title/company only; tailoring still works, weaker

    base = _docx_text(BASE_RESUME)
    resp = rq.post("https://api.anthropic.com/v1/messages",
                   headers={"x-api-key": key, "anthropic-version": "2023-06-01",
                            "content-type": "application/json"},
                   json={"model": "claude-sonnet-4-5", "max_tokens": 3000,
                         "system": TAILOR_SYSTEM,
                         "messages": [{"role": "user", "content":
                                       f"BASE RESUME:\n{base}\n\nJOB POSTING:\n{posting}"}]},
                   timeout=120)
    if resp.status_code != 200:
        return f"⛔ Anthropic API error {resp.status_code}: {resp.text[:200]}"
    tailored = "".join(b.get("text", "") for b in resp.json().get("content", []))

    safe = "".join(c if c.isalnum() else "_" for c in r["company"])[:30]
    out = RESUME_DIR / f"{safe}_row{job_number}.docx"
    _text_to_docx(tailored, out)
    r["resume_file"] = str(out)
    _save_tracker(rows)
    return (f"📄 Tailored resume for #{job_number} ({r['title']} @ {r['company']}):\n"
            f"`{out}`\n\nReview it line by line before applying. "
            f"Export to PDF via Word/Pages (File > Export). Real experience only; "
            f"verify nothing drifted.")


if __name__ == "__main__":
    mcp.run()
